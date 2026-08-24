"""
把「公共設施維護管理獎_評審委員問答準備手冊」轉成結構化知識庫
======================================================================
手冊本身就是高密度的問答對（含頁碼與具體數值），比一般文件更適合
以「一組問答＝一個知識單元」的方式儲存，讓 AI 答詢能精準命中，
不必依賴模糊檢索把整份文件切碎。

產出：webapp/data/review_handbook.json

用法：
    python scripts/build_review_handbook.py [來源.docx]
"""

from __future__ import annotations

import json
import os
import re
import sys

DEFAULT_SRC = os.path.join(
    os.path.expanduser('~'), 'Documents', '共設施維護管理獎評審委員名單',
    '公共設施維護管理獎_評審委員問答準備手冊_橫流溪.docx')

OUT_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                        'webapp', 'data', 'review_handbook.json')

# 各委員專屬問答表在 docx 中的出現順序（表 4 起，每位委員一張表）
REVIEWER_TABLE_ORDER = ['胡培中', '李振卿', '賴建宏', '王宜達', '張坤城']


def cell_text(cell) -> str:
    return re.sub(r'\s+\n', '\n', cell.text or '').strip()


def build(src: str) -> dict:
    from docx import Document
    doc = Document(src)

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = doc.tables

    data: dict = {
        'title': '公共工程金質獎「公共設施維護管理獎」評審委員問答準備手冊',
        'case': '橫流溪棲地連通性及周邊設施維護管理',
        'agency': '農業部林業及自然保育署臺中分署',
        'reviewers': [],
        'criteria': [],
        'platformModules': [],
        'qa': [],
        'briefingNotes': [],
        'overview': '',
    }

    if len(paragraphs) > 6:
        data['overview'] = paragraphs[6]

    # 表1：評分指標
    if len(tables) >= 1:
        for row in tables[0].rows[1:]:
            cells = [cell_text(c) for c in row.cells]
            if len(cells) >= 2 and cells[0]:
                data['criteria'].append({'指標': cells[0], '評審項目': cells[1]})

    # 表2：委員背景
    if len(tables) >= 2:
        for row in tables[1].rows[1:]:
            cells = [cell_text(c) for c in row.cells]
            if len(cells) >= 3 and cells[0]:
                data['reviewers'].append(
                    {'姓名': cells[0], '背景': cells[1], '關注重點': cells[2]})

    # 表3：平台功能模組
    if len(tables) >= 3:
        for row in tables[2].rows[1:]:
            cells = [cell_text(c) for c in row.cells]
            if len(cells) >= 2 and cells[0]:
                data['platformModules'].append(
                    {'模組': cells[0], '內容概要': cells[1]})

    # 表4 起：各委員的預期提問與建議回覆
    for index, table in enumerate(tables[3:]):
        reviewer = (REVIEWER_TABLE_ORDER[index]
                    if index < len(REVIEWER_TABLE_ORDER) else f'委員{index + 1}')
        for row in table.rows[1:]:
            cells = [cell_text(c) for c in row.cells]
            if len(cells) < 2 or not cells[0] or not cells[1]:
                continue
            question = cells[0]
            code = ''
            m = re.match(r'^(Q\d+)[.、．]\s*', question)
            if m:
                code = m.group(1)
                question = question[m.end():]
            data['qa'].append({
                'reviewer': reviewer,
                'code': code,
                'question': question,
                'answer': cells[1],
                'pages': sorted(set(re.findall(r'P\.\s*(\d+)', cells[0] + cells[1]))),
            })

    # 簡報前重點提醒（「五、」之後的段落）
    collecting = False
    for text in paragraphs:
        if text.startswith('五、') or '簡報前重點提醒' in text:
            collecting = True
            continue
        if collecting and len(text) > 20:
            data['briefingNotes'].append(text)

    # 各委員的敘述段（「四、N. XXX委員」後面兩段）
    for i, text in enumerate(paragraphs):
        m = re.match(r'^四、\s*\d+\.\s*(\S+?)委員', text)
        if not m:
            continue
        name = m.group(1)
        detail = [p for p in paragraphs[i + 1:i + 3] if len(p) > 10]
        for reviewer in data['reviewers']:
            if reviewer['姓名'] == name:
                if detail:
                    reviewer['職銜'] = detail[0]
                if len(detail) > 1:
                    reviewer['簡報建議'] = detail[1]
    return data


def main() -> int:
    src = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_SRC
    if not os.path.exists(src):
        print(f'找不到來源檔：{src}')
        return 1

    data = build(src)
    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    with open(OUT_PATH, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    print(f'已產出：{OUT_PATH}')
    print(f'  評分指標   {len(data["criteria"])} 項')
    print(f'  委員       {len(data["reviewers"])} 位')
    print(f'  平台模組   {len(data["platformModules"])} 個')
    print(f'  問答對     {len(data["qa"])} 組')
    print(f'  簡報提醒   {len(data["briefingNotes"])} 條')
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
