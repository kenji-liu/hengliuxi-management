from __future__ import annotations

import argparse
import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET


KEYWORDS = (
    "橫流溪", "物種", "尾數", "調查", "白甲魚", "石魚賓", "鬚鱲", "臺鰍",
    "吻鰕虎", "瘋鱨", "粗首馬口", "粗手馬口", "Zacco pachycephalus",
    "Opsariichthys pachycephalus",
)


def clean(value: object) -> str:
    text = "" if value is None else str(value)
    return re.sub(r"\s+", " ", text).strip()


def extract_docx(path: Path) -> list[str]:
    from docx import Document

    document = Document(path)
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = clean(paragraph.text)
        if text:
            lines.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        lines.append(f"[表格 {table_index}]")
        for row in table.rows:
            values = [clean(cell.text) for cell in row.cells]
            if any(values):
                lines.append("\t".join(values))
    return lines


def extract_pdf(path: Path) -> list[str]:
    import pdfplumber

    lines: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            lines.append(f"[PDF 第 {page_number} 頁]")
            text = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
            lines.extend(clean(line) for line in text.splitlines() if clean(line))
            for table_index, table in enumerate(page.extract_tables() or [], start=1):
                lines.append(f"[PDF 第 {page_number} 頁表格 {table_index}]")
                for row in table or []:
                    values = [clean(cell) for cell in (row or [])]
                    if any(values):
                        lines.append("\t".join(values))
    return lines


def extract_xlsx(path: Path) -> list[str]:
    from openpyxl import load_workbook

    workbook = load_workbook(path, data_only=True, read_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        lines.append(f"[工作表 {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            values = [clean(value) for value in row]
            while values and not values[-1]:
                values.pop()
            if any(values):
                lines.append("\t".join(values))
    return lines


def extract_ods(path: Path) -> list[str]:
    namespaces = {
        "office": "urn:oasis:names:tc:opendocument:xmlns:office:1.0",
        "table": "urn:oasis:names:tc:opendocument:xmlns:table:1.0",
        "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
    }
    repeat_col = f"{{{namespaces['table']}}}number-columns-repeated"
    repeat_row = f"{{{namespaces['table']}}}number-rows-repeated"
    office_value = f"{{{namespaces['office']}}}value"
    office_date = f"{{{namespaces['office']}}}date-value"
    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read("content.xml"))

    lines: list[str] = []
    for sheet in root.findall(".//table:table", namespaces):
        sheet_name = sheet.attrib.get(f"{{{namespaces['table']}}}name", "未命名")
        lines.append(f"[工作表 {sheet_name}]")
        for row in sheet.findall("table:table-row", namespaces):
            values: list[str] = []
            for cell in list(row):
                if not cell.tag.endswith(("table-cell", "covered-table-cell")):
                    continue
                paragraphs = [clean("".join(p.itertext())) for p in cell.findall(".//text:p", namespaces)]
                value = clean(" ".join(part for part in paragraphs if part))
                if not value:
                    value = clean(cell.attrib.get(office_date) or cell.attrib.get(office_value) or "")
                repeat = min(int(cell.attrib.get(repeat_col, "1")), 200)
                values.extend([value] * repeat)
            while values and not values[-1]:
                values.pop()
            if not any(values):
                continue
            row_repeat = min(int(row.attrib.get(repeat_row, "1")), 200)
            for _ in range(row_repeat):
                lines.append("\t".join(values))
    return lines


def context_matches(lines: list[str], radius: int = 3) -> list[dict[str, object]]:
    matched_indexes = [index for index, line in enumerate(lines) if any(keyword in line for keyword in KEYWORDS)]
    ranges: list[tuple[int, int]] = []
    for index in matched_indexes:
        start = max(0, index - radius)
        end = min(len(lines), index + radius + 1)
        if ranges and start <= ranges[-1][1]:
            ranges[-1] = (ranges[-1][0], max(ranges[-1][1], end))
        else:
            ranges.append((start, end))
    return [
        {"start_line": start + 1, "end_line": end, "text": lines[start:end]}
        for start, end in ranges
    ]


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source_dir", type=Path)
    parser.add_argument("output_dir", type=Path)
    args = parser.parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)

    extractors = {
        ".docx": extract_docx,
        ".pdf": extract_pdf,
        ".xlsx": extract_xlsx,
        ".ods": extract_ods,
    }
    report: list[dict[str, object]] = []
    for path in sorted(args.source_dir.iterdir(), key=lambda item: item.name):
        if not path.is_file():
            continue
        extractor = extractors.get(path.suffix.lower())
        entry: dict[str, object] = {"file": path.name, "type": path.suffix.lower()}
        if extractor is None:
            entry.update({"status": "requires_conversion", "line_count": 0, "matches": []})
            report.append(entry)
            continue
        try:
            lines = extractor(path)
            output_path = args.output_dir / f"{path.stem}.txt"
            output_path.write_text("\n".join(lines), encoding="utf-8")
            entry.update({
                "status": "extracted",
                "line_count": len(lines),
                "output": str(output_path),
                "matches": context_matches(lines),
            })
        except Exception as exc:
            entry.update({"status": "error", "error": f"{type(exc).__name__}: {exc}", "matches": []})
        report.append(entry)

    report_path = args.output_dir / "extraction_report.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps([
        {"file": item["file"], "status": item["status"], "line_count": item.get("line_count", 0), "match_blocks": len(item.get("matches", []))}
        for item in report
    ], ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
