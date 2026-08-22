#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
橫流溪 RAG 全資料庫索引器
索引三個根目錄的 PDF / DOCX / MD / JSON 檔案，
輸出與 rag_backend.py 相容的 vector_store.jsonl。

用法：
    python scripts/hlx_index_all.py            # 增量更新（只重索引變更的檔案）
    python scripts/hlx_index_all.py --force    # 強制重建整個向量庫
    python scripts/hlx_index_all.py --dry-run  # 只印出會索引哪些檔案，不實際建索引

相依：
    pip install sentence-transformers pypdf2 python-docx watchdog
"""

import argparse
import hashlib
import json
import logging
import re
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Any

# ── 路徑設定從統一 config 讀取 ──────────────────────────────────
SCRIPT_DIR = Path(__file__).parent
sys.path.insert(0, str(SCRIPT_DIR))

from hlx_rag_config import (
    SOURCE_ROOTS, VECTOR_STORE_FILE, METADATA_FILE, INDEX_MANIFEST,
    SUPPORTED_EXTS, EXCLUDE_DIRS, EXCLUDE_FILES, EXCLUDE_SUFFIXES,
    CATEGORY_RULES, EMBED_MODEL,
    CHUNK_SIZE, CHUNK_OVERLAP, MIN_CHUNK,
    PRIORITY_FILES,
)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("hlx-index")


# ═══════════════════════════════════════════════════════════════════
# 1. 嵌入：Jina AI API（雲端，免本機大模型）/ 本機 fallback
# ═══════════════════════════════════════════════════════════════════
import urllib.request
import urllib.error

JINA_API_KEY  = os.environ.get('JINA_API_KEY', '')
JINA_EMBED_URL = 'https://api.jina.ai/v1/embeddings'


def embed_texts_jina(texts: list[str], task: str = 'retrieval.passage') -> list[list[float]] | None:
    """呼叫 Jina AI API 取得向量。失敗回傳 None。"""
    if not texts:
        return []
    batch_size = 100
    all_vecs: list[list[float]] = []
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        payload = json.dumps({
            'model': EMBED_MODEL,            # 'jina-embeddings-v3'
            'input': batch,
            'task': task,
            'dimensions': 768,
        }).encode('utf-8')
        req = urllib.request.Request(
            JINA_EMBED_URL,
            data=payload,
            headers={
                'Authorization': f'Bearer {JINA_API_KEY}',
                'Content-Type': 'application/json',
                'Accept': 'application/json',
            },
            method='POST',
        )
        try:
            with urllib.request.urlopen(req, timeout=60) as resp:
                data = json.loads(resp.read())
            all_vecs.extend(item['embedding'] for item in data['data'])
        except urllib.error.HTTPError as e:
            log.error(f"Jina API HTTP {e.code}: {e.reason}")
            return None
        except Exception as e:
            log.error(f"Jina API 錯誤: {e}")
            return None
    return all_vecs


def embed_texts_local(texts: list[str]) -> list[list[float]] | None:
    """本機 sentence-transformers（備用）"""
    try:
        from sentence_transformers import SentenceTransformer
        model = SentenceTransformer(EMBED_MODEL)
        vecs = model.encode(texts, batch_size=32, show_progress_bar=False)
        return [v.tolist() for v in vecs]
    except ImportError:
        return None


def embed_texts(texts: list[str]) -> list[list[float]]:
    """自動選擇嵌入方式：優先 Jina API，fallback 本機模型"""
    if JINA_API_KEY:
        result = embed_texts_jina(texts)
        if result is not None:
            return result
        log.warning("Jina API 失敗，嘗試本機模型")
    result = embed_texts_local(texts)
    if result is None:
        log.error("無可用嵌入方式（設定 JINA_API_KEY 或安裝 sentence-transformers）")
        sys.exit(1)
    return result


# ═══════════════════════════════════════════════════════════════════
# 1b. Firebase Storage 上傳
# ═══════════════════════════════════════════════════════════════════

FIREBASE_STORAGE_BUCKET      = os.environ.get('FIREBASE_STORAGE_BUCKET', '')
FIREBASE_VECTOR_STORE_OBJECT = 'hlx/vector_store.jsonl'


def upload_to_firebase_storage(local_path: Path) -> bool:
    """將向量庫 JSONL 上傳至 Firebase Storage。需設定 FIREBASE_STORAGE_BUCKET。"""
    if not FIREBASE_STORAGE_BUCKET:
        log.info("FIREBASE_STORAGE_BUCKET 未設定，跳過上傳")
        return False

    # 優先使用 firebase-admin SDK（有服務帳號 JSON 時）
    try:
        import firebase_admin
        from firebase_admin import credentials, storage
        if not firebase_admin._apps:
            cred_path = os.environ.get('GOOGLE_APPLICATION_CREDENTIALS', '')
            if cred_path and Path(cred_path).exists():
                cred = credentials.Certificate(cred_path)
                firebase_admin.initialize_app(cred, {'storageBucket': FIREBASE_STORAGE_BUCKET})
            else:
                firebase_admin.initialize_app(options={'storageBucket': FIREBASE_STORAGE_BUCKET})
        bucket = storage.bucket()
        blob = bucket.blob(FIREBASE_VECTOR_STORE_OBJECT)
        blob.upload_from_filename(str(local_path), content_type='application/jsonl')
        log.info(f"✓ 上傳至 Firebase Storage: gs://{FIREBASE_STORAGE_BUCKET}/{FIREBASE_VECTOR_STORE_OBJECT}")
        # 設定為公開讀取（讓 Render 後端可以下載）
        blob.make_public()
        log.info(f"  公開 URL: {blob.public_url}")
        return True
    except ImportError:
        pass  # firebase-admin 未安裝，改用 REST API
    except Exception as e:
        log.error(f"firebase-admin 上傳失敗: {e}")

    # fallback：提示使用者手動上傳
    log.warning("firebase-admin SDK 未安裝（pip install firebase-admin）")
    log.warning(f"請手動上傳 {local_path} 至 Firebase Storage:")
    log.warning(f"  gs://{FIREBASE_STORAGE_BUCKET}/{FIREBASE_VECTOR_STORE_OBJECT}")
    return False


# ═══════════════════════════════════════════════════════════════════
# 2. 文字提取（PDF / DOCX / MD / JSON）
# ═══════════════════════════════════════════════════════════════════

def extract_pdf(path: Path) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(str(path))
        parts = []
        for page in reader.pages:
            t = page.extract_text() or ""
            if t.strip():
                parts.append(t)
        return "\n\n".join(parts)
    except Exception:
        try:
            from pypdf2 import PdfReader as PdfReader2
            reader2 = PdfReader2(str(path))
            parts = []
            for page in reader2.pages:
                t = page.extract_text() or ""
                if t.strip():
                    parts.append(t)
            return "\n\n".join(parts)
        except Exception as e:
            log.warning(f"  PDF提取失敗 {path.name}: {e}")
            return ""


def extract_docx(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
        parts = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        log.warning(f"  DOCX提取失敗 {path.name}: {e}")
        return ""


def extract_md(path: Path) -> str:
    try:
        text = path.read_text(encoding="utf-8", errors="ignore")
        # 移除 frontmatter
        text = re.sub(r"^---\s*\n.*?\n---\s*\n", "", text, flags=re.DOTALL)
        # 移除 Obsidian 特殊語法
        text = re.sub(r"\[\[([^\]]+)\]\]", r"\1", text)
        text = re.sub(r"!\[\[([^\]]+)\]\]", "", text)
        return text
    except Exception as e:
        log.warning(f"  MD提取失敗 {path.name}: {e}")
        return ""


def extract_json(path: Path) -> str:
    """專門處理 fish.js 等結構化資料"""
    try:
        text = path.read_text(encoding="utf-8", errors="ignore")
        # 如果是 fish.js 之類有 HLX_FISH_SURVEYS 的 JS 檔，提取資料陣列
        if "HLX_FISH_SURVEYS" in text or "label:" in text:
            # 把 label: '...' 提取出來
            labels = re.findall(r"label\s*:\s*'([^']+)'", text)
            # 再提取關鍵數值
            bai = re.findall(r"bai\s*:\s*(\d+)", text)
            if labels:
                pairs = []
                for i, lbl in enumerate(labels):
                    b = bai[i] if i < len(bai) else "?"
                    pairs.append(f"{lbl}: 白甲魚={b}")
                return "魚類調查紀錄：\n" + "\n".join(pairs)
        # 一般 JSON 檔：直接序列化成可讀文字
        data = json.loads(text)
        return json.dumps(data, ensure_ascii=False, indent=2)[:8000]
    except Exception:
        try:
            return path.read_text(encoding="utf-8", errors="ignore")[:8000]
        except Exception as e:
            log.warning(f"  JSON提取失敗 {path.name}: {e}")
            return ""


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf(path)
    elif ext in (".docx", ".doc"):
        return extract_docx(path)
    elif ext in (".md", ".txt"):
        return extract_md(path)
    elif ext in (".json",):
        return extract_json(path)
    else:
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""


# ═══════════════════════════════════════════════════════════════════
# 3. 分塊
# ═══════════════════════════════════════════════════════════════════

def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """以段落優先的分塊策略"""
    # 先以多個空行分段
    paragraphs = re.split(r"\n{2,}", text)
    chunks: list[str] = []
    buf = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(buf) + len(para) < size:
            buf = (buf + "\n\n" + para).strip()
        else:
            if buf and len(buf) >= MIN_CHUNK:
                chunks.append(buf)
            # 若單段過長，強制切割
            if len(para) > size:
                for i in range(0, len(para), size - overlap):
                    sub = para[i:i + size]
                    if len(sub) >= MIN_CHUNK:
                        chunks.append(sub)
                buf = para[-(overlap):]
            else:
                buf = para

    if buf and len(buf) >= MIN_CHUNK:
        chunks.append(buf)

    return chunks


# ═══════════════════════════════════════════════════════════════════
# 4. 分類判定
# ═══════════════════════════════════════════════════════════════════

def detect_category(path: Path, text: str) -> str:
    full = (str(path) + " " + text[:2000]).lower()
    scores: dict[str, int] = {}
    for cat, keywords in CATEGORY_RULES:
        score = sum(1 for kw in keywords if kw.lower() in full)
        if score:
            scores[cat] = score

    if not scores:
        # 依目錄猜測
        p = str(path)
        if "魚類" in p or "fish" in p.lower():
            return "生態資料庫"
        if "巡查" in p or "inspection" in p.lower():
            return "巡查資料管理"
        if "維護" in p or "設施" in p:
            return "工程設施管理"
        return "工程設計書架"

    return max(scores, key=lambda k: scores[k])


# ═══════════════════════════════════════════════════════════════════
# 5. 檔案掃描
# ═══════════════════════════════════════════════════════════════════

def should_skip(path: Path) -> bool:
    # 跳過隱藏檔
    if any(part.startswith(".") for part in path.parts):
        return True
    # 跳過排除目錄
    for part in path.parts:
        if part in EXCLUDE_DIRS:
            return True
    # 跳過排除檔名
    if path.name in EXCLUDE_FILES:
        return True
    # 跳過不支援的附檔名
    if path.suffix.lower() in EXCLUDE_SUFFIXES:
        return True
    # 跳過太小的檔案（< 100 bytes）
    try:
        if path.stat().st_size < 100:
            return True
    except Exception:
        return True
    return False


def gather_files() -> list[Path]:
    """收集所有根目錄下的可索引檔案，依優先清單排序"""
    all_files: list[Path] = []

    for root in SOURCE_ROOTS:
        if not root.exists():
            log.warning(f"目錄不存在，跳過：{root}")
            continue
        log.info(f"掃描目錄：{root}")
        try:
            for path in root.rglob("*"):
                try:
                    if not path.is_file():
                        continue
                except OSError:
                    continue  # WinError 1920: junction/symlink 無法存取
                if path.suffix.lower() not in SUPPORTED_EXTS:
                    continue
                if should_skip(path):
                    continue
                all_files.append(path)
        except (OSError, PermissionError) as e:
            log.warning(f"掃描 {root} 時發生錯誤（部分跳過）：{e}")

    # 優先排序
    def priority_key(p: Path) -> int:
        name = str(p)
        for i, keyword in enumerate(PRIORITY_FILES):
            if keyword in name:
                return i
        return len(PRIORITY_FILES)

    all_files.sort(key=priority_key)
    log.info(f"找到 {len(all_files)} 個可索引檔案")
    return all_files


# ═══════════════════════════════════════════════════════════════════
# 6. manifest（增量更新用）
# ═══════════════════════════════════════════════════════════════════

def file_hash(path: Path) -> str:
    h = hashlib.md5()
    try:
        h.update(path.read_bytes())
    except Exception:
        h.update(str(path.stat().st_mtime).encode())
    return h.hexdigest()


def load_manifest() -> dict[str, Any]:
    if INDEX_MANIFEST.exists():
        try:
            return json.loads(INDEX_MANIFEST.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {}


def save_manifest(manifest: dict[str, Any]):
    INDEX_MANIFEST.parent.mkdir(parents=True, exist_ok=True)
    INDEX_MANIFEST.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


# ═══════════════════════════════════════════════════════════════════
# 7. 主索引邏輯
# ═══════════════════════════════════════════════════════════════════

def index_all(force: bool = False, dry_run: bool = False):
    files = gather_files()
    manifest = {} if force else load_manifest()

    # 找出需要重新索引的檔案
    to_index: list[Path] = []
    unchanged: list[Path] = []
    for f in files:
        h = file_hash(f)
        if force or manifest.get(str(f)) != h:
            to_index.append(f)
        else:
            unchanged.append(f)

    log.info(f"需重索引：{len(to_index)} 個；未變更（直接保留）：{len(unchanged)} 個")

    if dry_run:
        print("\n=== DRY RUN：將要索引的檔案 ===")
        for f in to_index:
            print(f"  {f}")
        print(f"\n共 {len(to_index)} 個檔案將被索引")
        return

    if not to_index and not force:
        log.info("所有檔案均未變更，無需重新索引")
        return

    # 確認嵌入方式
    if JINA_API_KEY:
        log.info(f"嵌入模式：Jina AI API（{EMBED_MODEL}）")
    else:
        log.info(f"嵌入模式：本機 sentence-transformers（{EMBED_MODEL}）")

    # 如果是增量更新，先讀取現有向量庫（保留未變更的條目）
    existing_entries: list[dict] = []
    if not force and VECTOR_STORE_FILE.exists():
        unchanged_strs = {str(f) for f in unchanged}
        log.info("讀取現有向量庫（保留未變更條目）…")
        with open(VECTOR_STORE_FILE, encoding="utf-8") as fh:
            for line in fh:
                line = line.strip()
                if not line:
                    continue
                try:
                    entry = json.loads(line)
                    if entry.get("source", "") in unchanged_strs:
                        existing_entries.append(entry)
                except Exception:
                    pass
        log.info(f"保留 {len(existing_entries)} 個現有向量條目")

    # 開始處理新檔案
    new_entries: list[dict] = []
    new_manifest: dict[str, str] = {str(f): manifest.get(str(f), "") for f in unchanged}

    for i, path in enumerate(to_index, 1):
        log.info(f"[{i}/{len(to_index)}] {path.name}")
        text = extract_text(path)
        if not text.strip():
            log.warning(f"  → 無法提取文字，跳過")
            continue

        category = detect_category(path, text)
        chunks = chunk_text(text)
        if not chunks:
            log.warning(f"  → 分塊為空，跳過")
            continue

        log.info(f"  → 分類={category}，{len(chunks)} 個塊")

        # 批次向量化
        try:
            vectors = embed_texts(chunks)
        except Exception as e:
            log.error(f"  → 向量化失敗：{e}")
            continue

        for j, (chunk, vector) in enumerate(zip(chunks, vectors)):
            entry = {
                "text":        chunk,
                "vector":      vector,
                "source":      str(path),
                "source_path": str(path),   # rag_backend.py is_target_topic_doc() reads this field
                "source_file": path.name,   # normalize_source_name() reads this field
                "filename":    path.name,
                "category":    category,
                "chunk_id":    j,
                "indexed_at":  datetime.now().isoformat(timespec="seconds"),
            }
            # 補充 metadata
            if path.suffix.lower() == ".pdf":
                entry["filetype"] = "pdf"
            elif path.suffix.lower() in (".docx", ".doc"):
                entry["filetype"] = "docx"
            elif path.suffix.lower() == ".md":
                entry["filetype"] = "markdown"
            else:
                entry["filetype"] = path.suffix.lstrip(".")
            new_entries.append(entry)

        new_manifest[str(path)] = file_hash(path)

    # 寫出向量庫
    all_entries = existing_entries + new_entries
    VECTOR_STORE_FILE.parent.mkdir(parents=True, exist_ok=True)
    log.info(f"寫出向量庫（共 {len(all_entries)} 個條目）→ {VECTOR_STORE_FILE}")
    with open(VECTOR_STORE_FILE, "w", encoding="utf-8") as fh:
        for entry in all_entries:
            fh.write(json.dumps(entry, ensure_ascii=False) + "\n")

    # 更新 manifest
    save_manifest(new_manifest)

    # 更新 metadata 摘要
    categories: dict[str, int] = {}
    for e in all_entries:
        cat = e.get("category", "未分類")
        categories[cat] = categories.get(cat, 0) + 1

    meta_summary = {
        "total_chunks": len(all_entries),
        "total_files":  len(new_manifest),
        "categories":   categories,
        "embed_model":  EMBED_MODEL,
        "last_indexed": datetime.now().isoformat(timespec="seconds"),
        "source_roots": [str(r) for r in SOURCE_ROOTS],
    }
    METADATA_FILE.write_text(
        json.dumps(meta_summary, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    log.info("索引完成！")
    log.info(f"  各分類統計：{categories}")
    log.info(f"  向量庫位置：{VECTOR_STORE_FILE}")

    # ── 上傳至 Firebase Storage（雲端部署用）────────────────────
    log.info("嘗試上傳向量庫至 Firebase Storage…")
    uploaded = upload_to_firebase_storage(VECTOR_STORE_FILE)
    if uploaded:
        log.info("✓ Firebase Storage 上傳完成，Render 後端可自動下載最新向量庫")
        log.info("  在網頁 AI 問答管理面板按「重新載入知識庫」即生效")
    else:
        log.info("  跳過 Firebase 上傳（本機模式或未設定 FIREBASE_STORAGE_BUCKET）")


# ═══════════════════════════════════════════════════════════════════
# 8. CLI
# ═══════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="橫流溪 RAG 全資料庫索引器")
    parser.add_argument("--force",   action="store_true", help="強制重建整個向量庫")
    parser.add_argument("--dry-run", action="store_true", help="只列出要索引的檔案，不實際建索引")
    args = parser.parse_args()

    index_all(force=args.force, dry_run=args.dry_run)
